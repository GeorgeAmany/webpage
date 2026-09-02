#!/usr/bin/env python3
"""Generate synced George Amany CV Word document from evidence-based content."""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from pathlib import Path

ROOT = Path(__file__).resolve().parent.parent
OUTPUT = ROOT / "assets" / "George_Amany_CV_2026.docx"

CONTACT_LINKS = {
    "email": "mailto:georgeamany5@gmail.com",
    "linkedin": "https://www.linkedin.com/in/george-amany-53b148219/",
    "github": "https://github.com/GeorgeAmany",
}

PROJECT_LINKS = {
    "Noor Institute": [
        ("Google Play", "https://play.google.com/store/apps/details?id=com.softera.noor_academy"),
        ("App Store", "https://apps.apple.com/eg/app/noor-institute/id6463731504"),
    ],
    "Schupply": [
        ("Google Play", "https://play.google.com/store/apps/details?id=com.schupply.app"),
    ],
    "Binge": [
        ("Google Play", "https://play.google.com/store/apps/details?id=com.jigsaw.binge&hl=ar"),
    ],
    "Al Wefaq Foods": [
        ("Google Play", "https://play.google.com/store/apps/details?id=com.alwefaqfoods.app"),
    ],
    "Horse Time": [
        ("Google Play", "https://play.google.com/store/apps/details?id=com.horsetime.user&hl=en"),
        ("App Store", "https://apps.apple.com/eg/app/horse-time/id6758008343"),
    ],
    "Al Rassi": [
        ("Google Play", "https://play.google.com/store/apps/details?id=com.appsbunches.alrassiapp"),
        ("App Store", "https://apps.apple.com/eg/app/%D8%A7%D9%84%D8%B1%D8%B3%D9%8A-%D9%84%D9%84%D8%A7%D8%B4%D8%AC%D8%A7%D8%B1-%D8%A7%D9%84%D8%B5%D9%86%D8%A7%D8%B9%D9%8A%D8%A9/id6736564188"),
    ],
    "animated_contact_us": [
        ("pub.dev", "https://pub.dev/packages/animated_contact_us"),
        ("GitHub", "https://github.com/GeorgeAmany/animated_contact_us"),
    ],
    "liquid_wave_indicator": [
        ("pub.dev", "https://pub.dev/packages/liquid_wave_indicator"),
        ("GitHub", "https://github.com/GeorgeAmany/liquid_wave_indicator"),
    ],
}


def add_hyperlink(paragraph, url, text):
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0F766E")
    r_pr.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(underline)
    run.append(r_pr)
    text_elem = OxmlElement("w:t")
    text_elem.text = text
    run.append(text_elem)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_project_links(paragraph, links):
    for i, (label, url) in enumerate(links):
        if i:
            paragraph.add_run(" · ")
        add_hyperlink(paragraph, url, label)

doc = Document()
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)

# --- Header ---
name = doc.add_paragraph()
name.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = name.add_run("George Amany")
run.bold = True
run.font.size = Pt(22)

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("Mid Flutter Developer | Mobile & Desktop Cross-Platform Engineer")
run.font.size = Pt(12)

contact = doc.add_paragraph()
contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
contact.add_run("Cairo, Egypt | ")
add_hyperlink(contact, CONTACT_LINKS["email"], "georgeamany5@gmail.com")
contact.add_run(" | +20 127 003 7845 | ")
add_hyperlink(contact, CONTACT_LINKS["linkedin"], "LinkedIn")
contact.add_run(" | ")
add_hyperlink(contact, CONTACT_LINKS["github"], "GitHub")

doc.add_paragraph()

# --- Professional Summary ---
doc.add_heading("Professional Summary", level=1)
doc.add_paragraph(
    "Flutter engineer with a Bachelor's degree in Computer Science and production experience "
    "building and shipping mobile and desktop applications across education, e-commerce, food "
    "services, and influencer marketing. At Neural Design Studios, delivers feature-first Clean "
    "Architecture apps using BLoC/Cubit, Dio/Retrofit REST integrations, Firebase, Hive/Floor "
    "local storage, and cross-platform desktop tooling (macOS/Windows). Published multiple apps "
    "to Google Play and the App Store, authored open-source packages on pub.dev, and set up "
    "GitHub Actions CI for automated APK builds. Grew quickly under senior mentorship and now "
    "mentors interns and junior developers on Flutter, Clean Architecture, and code-quality standards."
)

# --- Core Skills ---
doc.add_heading("Core Skills", level=1)
skills = [
    "Flutter & Dart", "Clean Architecture", "BLoC / Cubit State Management", "GetIt / Injectable DI",
    "Dio & Retrofit REST APIs", "Firebase (Auth, Firestore, Messaging, Crashlytics, Analytics)",
    "Hive, Floor & SQLite Local Storage", "GoRouter Navigation", "Freezed & Dartz",
    "Stripe & Multi-Gateway Payments", "Pusher Real-Time", "Flutter Desktop (macOS/Windows)",
    "Git & GitHub Actions CI/CD", "Easy Localization", "Postman", "Clean Code & Code Review",
    "Mentoring Interns & Junior Developers",
]
p = doc.add_paragraph(", ".join(skills))

# --- Experience ---
doc.add_heading("Professional Experience", level=1)

doc.add_heading("Flutter Developer — Neural Design Studios", level=2)
doc.add_paragraph("Full-time | Cairo, Egypt | Jun 2024 – Present")
bullets = [
    "Shipped production Flutter apps across education, e-commerce, food, and influencer marketing domains with feature-first Clean Architecture (data/domain/presentation layers, repository pattern, use cases).",
    "Led development on Fuse influencer marketing platform: dual-app architecture for agencies and influencers with Cubit, Injectable DI, Freezed models, Dio networking, Firebase Auth/Realtime DB, and Storyly integration.",
    "Built Taleem student and employee apps for Islamic education: REST v2 API integration, Quran local SQLite databases, attendance/memorization/evaluation modules, and GitHub Actions CI pipeline for automated debug APK builds.",
    "Enhanced Noor Institute app: Stripe payment flow, Pusher real-time chat, Floor local DB, Retrofit API layer, Syncfusion calendar, and Firebase push notifications.",
    "Delivered cross-platform desktop apps: Zahran POS (offline invoice sync, PDF printing, Hive storage) and Proposal Desktop (PDF document builder with flutter_bloc and window_manager).",
    "Integrated REST APIs, Firebase services, Google/Apple Sign-In, Google Maps, and payment gateways (Stripe, Amazon Payment Services) across Schupply, Zahran, Binge, and Al Wefaq Foods apps.",
    "Published and maintained open-source Flutter packages on pub.dev (animated_contact_us, liquid_wave_indicator).",
    "Mentor interns and junior developers — onboarding them onto Flutter, Clean Architecture, and code-review and code-quality practices, and supporting them in shipping their first production features.",
    "Applied maintainability practices: dependency injection, typed API clients, localization (easy_localization), structured error handling, and consistent feature module boundaries.",
]
for b in bullets:
    doc.add_paragraph(b, style="List Bullet")

doc.add_heading("Flutter Developer — Neural Design Studios", level=2)
doc.add_paragraph("Part-time | Cairo, Egypt | Feb 2024 – Jun 2024")
for b in [
    "Collaborated with cross-functional teams to deliver responsive, production-ready Flutter applications.",
    "Implemented UI performance optimizations and scalable state management patterns.",
]:
    doc.add_paragraph(b, style="List Bullet")

# --- Key Projects ---
doc.add_heading("Key Projects", level=1)

projects = [
    ("Noor Institute", "Education / Institute Management",
     "Production app for students, teachers, courses, homework, schedules, invoices, and chat. "
     "Clean Architecture with BLoC, Retrofit/Dio, Floor DB, Stripe payments, Pusher real-time, Firebase push.",
     "Flutter, BLoC, GetIt, Dio, Retrofit, Floor, Stripe, Pusher, Firebase, Freezed, Easy Localization",
     PROJECT_LINKS["Noor Institute"]),
    ("Fuse", "Influencer Marketing",
     "Dual-role platform for agencies and influencers: profiles, content posting, service requests, stories.",
     "Flutter, Cubit, Injectable, Freezed, Dio, Hive, Firebase Auth/DB, Storyly, GoRouter",
     []),
    ("Taleem (Student & Employees)", "Islamic Education",
     "Student app (grades, absence, certificates, Quran mushaf) and staff app (halaqas, attendance, memorization, evaluations). GitHub Actions CI for APK builds.",
     "Flutter, BLoC, Dio, Hive, SQLite, Firebase Messaging, Fl Chart, Table Calendar",
     []),
    ("Schupply", "School Supplies E-Commerce",
     "School/grade package ordering with REST API, Hive caching, Google/Apple auth, deep links.",
     "Flutter, BLoC, Dio, Hive, GoRouter, Firebase Messaging",
     PROJECT_LINKS["Schupply"]),
    ("Zahran (Mobile & Desktop POS)", "E-Commerce & Point of Sale",
     "Mobile e-commerce app and desktop POS with offline invoice sync, PDF export/printing, secure storage.",
     "Flutter, BLoC, Dio, Hive, Window Manager, Printing, PDF, Connectivity Plus",
     []),
    ("Binge", "Food Subscription & Meal Ordering",
     "Dishes, cart, subscriptions, wallet, loyalty program. Firebase Analytics/Crashlytics, Amazon Payment Services.",
     "Flutter, BLoC, Dio, Hive, Firebase, Google Maps, Clarity Analytics",
     PROJECT_LINKS["Binge"]),
    ("Al Wefaq Foods", "Food Brand Mobile",
     "WebView wrapper with Firebase/OneSignal push, in-app messaging, QR scanner.",
     "Flutter, InAppWebView, MobX, Firebase, OneSignal",
     PROJECT_LINKS["Al Wefaq Foods"]),
    ("Proposal Desktop", "Business Proposal Builder",
     "Desktop PDF proposal generator with editor, service packages, company info, payment methods.",
     "Flutter Desktop, BLoC, Dio, Hive, PDF, Printing, Window Manager",
     []),
    ("Horse Time", "On-Demand Service Booking",
     "User booking app with chat, payments, maps, and multi-gateway payment integration.",
     "Flutter, MobX, Firebase, Google Maps, Stripe/Razorpay/PayPal",
     PROJECT_LINKS["Horse Time"]),
    ("animated_contact_us", "Open Source (pub.dev)",
     "Published Flutter UI package for animated contact widgets.",
     "Flutter, Font Awesome, URL Launcher",
     PROJECT_LINKS["animated_contact_us"]),
    ("liquid_wave_indicator", "Open Source (pub.dev)",
     "Published Flutter UI package for liquid wave progress indicators.",
     "Flutter",
     PROJECT_LINKS["liquid_wave_indicator"]),
    ("Al Rassi", "Industrial / Manufacturing",
     "Published production mobile application.",
     "Flutter",
     PROJECT_LINKS["Al Rassi"]),
]

for name, domain, desc, tech, links in projects:
    doc.add_heading(name, level=2)
    doc.add_paragraph(f"Domain: {domain}")
    doc.add_paragraph(desc)
    doc.add_paragraph(f"Technologies: {tech}")
    if links:
        links_p = doc.add_paragraph()
        links_p.add_run("Links: ")
        add_project_links(links_p, links)

# --- Education ---
doc.add_heading("Education", level=1)
doc.add_paragraph("Bachelor's Degree, Computer Science — Modern Academy Maadi, Cairo (2022)")

# --- Courses & Activities ---
doc.add_heading("Courses", level=1)
for c in [
    "Full Flutter Diploma (100 hours) — Array",
    "Flutter & Dart — Udemy",
    "Cyber Security Diploma (150 hours) — Instant",
]:
    doc.add_paragraph(c, style="List Bullet")

doc.add_heading("Community & Activities", level=1)
for a in [
    "Flutter workshop — Alalmiya Alhura, El Mansoura",
    "Flutter Forward Extended Cloud Egypt '23 — GDG Cloud Egypt",
    "Online Flutter workshop — Machinfy",
    "Penetration Tester Training (1 month) — Instant",
]:
    doc.add_paragraph(a, style="List Bullet")

doc.add_heading("Languages", level=1)
doc.add_paragraph("Arabic (Native) | English (Professional Working Proficiency)")

doc.save(OUTPUT)
print(f"Saved: {OUTPUT}")
